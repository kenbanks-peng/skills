#!/usr/bin/env python3
"""Generate a Word (.docx) daily summary of jobs recorded today.

Usage: python daily_report.py <jobs.db> <output.docx> [--date YYYY-MM-DD]
                                                      [--search-spec <path>]

The output file should end in ``.docx``. If ``--search-spec`` is provided,
track labels and display order are read from that file. Tracks present in
the DB but absent from the spec are still rendered, sorted alphabetically
after the spec-defined ones, with the raw id as label. The bare slug
``unknown`` and SQL NULL both render as "Unclassified".

Each match row includes the posting's own job ID when the source board
exposed one (``source_job_id``), falling back to the internal DB row id in
parentheses so rows can always be cross-referenced against ``jobs.db``.

If --date is omitted, uses today. Output file's parent dir is created if
needed.

Dependency: python-docx. Install with:
    pip install python-docx --break-system-packages
"""

import argparse
import sqlite3
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "python-docx is required. Install with:\n"
        "    pip install python-docx --break-system-packages\n"
    )
    sys.exit(2)

# Re-use the search-spec parser. It lives next to this script.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from parse_search_spec import parse_search_spec  # noqa: E402


# --------------------------------------------------------------------------- #
# Data loading
# --------------------------------------------------------------------------- #


def load_track_metadata(spec_path: Path | None) -> tuple[dict, list]:
    """Return (labels, order) for tracks. Falls back to empty dicts if no spec."""
    labels: dict[str | None, str] = {"unknown": "Unclassified", None: "Unclassified"}
    order: list[str | None] = []
    if spec_path is None:
        return labels, order
    spec = parse_search_spec(spec_path.read_text(encoding="utf-8"))
    for track in spec.get("track") or []:
        tid = track["id"]
        labels[tid] = track.get("label") or tid
        order.append(tid)
    return labels, order


def fetch_rows(conn: sqlite3.Connection, target_date: str):
    return conn.execute(
        """SELECT id, job_title, company, location, job_url, source_url,
                  salary_raw, keyword_alignment, match_score, track, state,
                  source_job_id, source
           FROM jobs
           WHERE date_found = ?
           ORDER BY match_score DESC NULLS LAST, id DESC""",
        (target_date,),
    ).fetchall()


# --------------------------------------------------------------------------- #
# Row helpers
# --------------------------------------------------------------------------- #


def _row_track(row) -> str | None:
    return row[9]


def _label_for(track: str | None, labels: dict) -> str:
    if track in labels:
        return labels[track]
    return track if track else "Unclassified"


def _ordered_tracks(by_track: dict, spec_order: list) -> list:
    """Spec-defined tracks first (in spec order, only if present), then any
    extras the DB has but the spec doesn't, sorted alphabetically. ``unknown``
    and None always sort last."""
    seen: set = set()
    out: list = []
    for tid in spec_order:
        if tid in by_track:
            out.append(tid)
            seen.add(tid)
    extras = [
        t for t in by_track.keys() if t not in seen and t not in ("unknown", None)
    ]
    out.extend(sorted(extras, key=lambda x: (x is None, x or "")))
    for tail in ("unknown", None):
        if tail in by_track and tail not in out:
            out.append(tail)
    return out


def _format_job_id(db_id, source_job_id) -> str:
    """Display the posted (source) job ID when available, otherwise the DB id.

    Examples:
        source_job_id="4127884412", db_id=57  -> "4127884412 (#57)"
        source_job_id=None,        db_id=57  -> "#57"
    """
    sjid = (source_job_id or "").strip() if source_job_id is not None else ""
    if sjid:
        return f"{sjid} (#{db_id})"
    return f"#{db_id}"


# --------------------------------------------------------------------------- #
# docx rendering
# --------------------------------------------------------------------------- #

TABLE_HEADERS = [
    "Job ID",
    "Score",
    "Title",
    "Company",
    "Location",
    "Salary",
    "Alignment",
]

HEADER_BG = "1F2937"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(
    cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int = 9
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text if text is not None else "")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _add_hyperlink(paragraph, url: str, text: str, size: int = 9) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))  # half-points
    r_pr.append(sz)

    new_run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text or "(untitled)"
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _write_track_table(doc, rows) -> None:
    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header row
    hdr = table.rows[0].cells
    for i, header in enumerate(TABLE_HEADERS):
        _set_cell_text(hdr[i], header, bold=True, color=HEADER_FG, size=9)
        _shade_cell(hdr[i], HEADER_BG)

    # Body rows
    for (
        db_id,
        title,
        company,
        location,
        job_url,
        _source_url,
        salary_raw,
        alignment,
        score,
        _track,
        _state,
        source_job_id,
        _source,
    ) in rows:
        cells = table.add_row().cells

        _set_cell_text(cells[0], _format_job_id(db_id, source_job_id), size=9)
        _set_cell_text(
            cells[1],
            str(score) if score is not None else "—",
            size=9,
        )

        # Title is a hyperlink to the job URL.
        cells[2].text = ""
        p = cells[2].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if job_url:
            _add_hyperlink(p, job_url, title or "(untitled)", size=9)
        else:
            run = p.add_run(title or "(untitled)")
            run.font.size = Pt(9)

        _set_cell_text(cells[3], company or "—", size=9)
        _set_cell_text(cells[4], location or "—", size=9)
        _set_cell_text(cells[5], salary_raw or "—", size=9)
        _set_cell_text(cells[6], alignment or "—", size=9)

    # Column widths (approximate; Word will adjust for content).
    widths = [
        Inches(1.2),
        Inches(0.5),
        Inches(2.0),
        Inches(1.3),
        Inches(1.3),
        Inches(1.1),
        Inches(1.6),
    ]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]


def build_report(
    target_date: str,
    rows,
    labels: dict,
    spec_order: list,
    out_path: Path,
) -> None:
    doc = Document()

    # Page margins — a little tighter than the Word default so the match
    # tables fit without wrapping.
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Normal style: a more readable body font at 10pt.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    title = doc.add_heading(f"Job search report — {target_date}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    summary = doc.add_paragraph()
    srun = summary.add_run(f"New matches recorded today: {len(rows)}")
    srun.bold = True

    if not rows:
        doc.add_paragraph("No new matches today. Run completed successfully.")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return

    # Per-track grouping
    by_track: dict[str | None, list] = {}
    for row in rows:
        by_track.setdefault(_row_track(row), []).append(row)

    track_order = _ordered_tracks(by_track, spec_order)

    # By-track totals
    doc.add_heading("By track", level=2)
    for track in track_order:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(_label_for(track, labels))
        run.bold = True
        p.add_run(f" ({track or 'null'}): {len(by_track[track])}")

    # By-source totals
    by_source: dict[str, int] = {}
    for row in rows:
        source_url = row[5]
        key = source_url or "(unknown source)"
        by_source[key] = by_source.get(key, 0) + 1

    doc.add_heading("By source", level=2)
    for src, count in sorted(by_source.items(), key=lambda kv: -kv[1]):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{count} — {src}")

    # One section per track with a match table
    for track in track_order:
        doc.add_heading(f"{_label_for(track, labels)} — new matches", level=2)
        _write_track_table(doc, by_track[track])
        doc.add_paragraph("")  # spacer between tables

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("db")
    parser.add_argument("out", help="Output .docx path")
    parser.add_argument("--date", default=date.today().isoformat())
    parser.add_argument(
        "--search-spec",
        default=None,
        help="Optional path to job-search-spec.toml for track labels/order",
    )
    args = parser.parse_args()

    db_path = Path(args.db).expanduser()
    out_path = Path(args.out).expanduser()
    if out_path.suffix.lower() != ".docx":
        sys.stderr.write(
            f"warning: output path {out_path} does not end in .docx — "
            "daily reports are Word documents\n"
        )
    spec_path = Path(args.search_spec).expanduser() if args.search_spec else None

    labels, spec_order = load_track_metadata(spec_path)

    conn = sqlite3.connect(db_path)
    try:
        rows = fetch_rows(conn, args.date)
    finally:
        conn.close()

    build_report(args.date, rows, labels, spec_order, out_path)
    print(f"wrote {out_path} ({len(rows)} rows)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
